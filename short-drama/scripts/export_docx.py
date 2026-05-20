#!/usr/bin/env python3
"""短剧剧本 Markdown → Word 导出脚本（纯 python-docx 版）

用法: python3 scripts/export_docx.py <输入.md> <输出.docx>

功能:
  1. 检测 python-docx / lxml 是否安装，未安装时给出安装指令
  2. 逐行解析剧本 Markdown（# / ## / ### / △ / 对白 / 括号 / 分隔线）
  3. 用 python-docx 直接生成 .docx，精确控制中文字体（w:eastAsia）和间距
  4. 参照专业剧本排版：宋体/黑体 12pt，全文加粗，1.5倍行距

注意: 第三个位置参数（reference-doc 路径）已不再使用，保留接口兼容性。
"""

import re
import sys
from pathlib import Path


# ─────────────────────── 依赖检测 ───────────────────────

def check_deps() -> bool:
    missing = []
    try:
        import docx  # noqa
    except ImportError:
        missing.append("python-docx")
    try:
        from lxml import etree  # noqa
    except ImportError:
        missing.append("lxml")

    if missing:
        pkgs = " ".join(missing)
        print(f"[错误] 缺少依赖，请安装: pip3 install {pkgs}", file=sys.stderr)
        return False
    return True


# ─────────────────────── 文档初始化 ───────────────────────

def _set_east_asia(rpr_elem, font_cn: str):
    """在 rPr XML 元素上设置 w:eastAsia 中文字体属性"""
    from docx.oxml.ns import qn
    from lxml import etree

    rFonts = rpr_elem.find(qn("w:rFonts"))
    if rFonts is None:
        rFonts = etree.SubElement(rpr_elem, qn("w:rFonts"))
    rFonts.set(qn("w:eastAsia"), font_cn)


def init_document():
    """创建文档，设置 A4 页面 + 全局 Normal 样式"""
    from docx import Document
    from docx.shared import Pt, Cm

    doc = Document()

    # A4 页面，标准页边距
    sec = doc.sections[0]
    sec.page_width  = Cm(21.0)
    sec.page_height = Cm(29.7)
    sec.top_margin    = Cm(2.54)
    sec.bottom_margin = Cm(2.54)
    sec.left_margin   = Cm(3.17)
    sec.right_margin  = Cm(3.17)

    # Normal 基础样式：宋体 12pt 加粗，1.5倍行距，6pt 段后间距
    normal = doc.styles["Normal"]
    normal.font.size  = Pt(12)
    normal.font.bold  = True
    normal.font.name  = "Times New Roman"   # 英文/Latin 字体
    normal.paragraph_format.line_spacing  = 1.5
    normal.paragraph_format.space_before  = Pt(0)
    normal.paragraph_format.space_after   = Pt(6)
    _set_east_asia(normal.element.get_or_add_rPr(), "宋体")

    return doc


# ─────────────────────── Run 辅助 ───────────────────────

def _make_run(para, text: str, bold: bool = True, size_pt: int = 12,
              font_en: str = "Times New Roman", font_cn: str = "宋体"):
    """添加一个 run，同时设置中英文字体"""
    from docx.shared import Pt
    from docx.oxml.ns import qn
    from lxml import etree

    run = para.add_run(text)
    run.bold       = bold
    run.font.size  = Pt(size_pt)
    run.font.name  = font_en

    rPr = run._r.get_or_add_rPr()
    rFonts = rPr.find(qn("w:rFonts"))
    if rFonts is None:
        rFonts = etree.SubElement(rPr, qn("w:rFonts"))
    rFonts.set(qn("w:eastAsia"), font_cn)
    return run


def _add_para(doc, text: str, *,
              bold: bool = True, size_pt: int = 12,
              font_en: str = "Times New Roman", font_cn: str = "宋体",
              space_before: float = 0, space_after: float = 6,
              alignment=None):
    """
    添加段落，自动解析内联 **粗体** 标记。
    ** 包裹的部分强制加粗，其余部分使用 bold 参数值。
    """
    from docx.shared import Pt

    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(space_before)
    p.paragraph_format.space_after  = Pt(space_after)
    if alignment is not None:
        p.paragraph_format.alignment = alignment

    # 拆分内联 **bold** 标记（非贪婪，不跨行）
    segments = re.split(r"(\*\*[^*\n]+\*\*)", text)
    for seg in segments:
        if not seg:
            continue
        if seg.startswith("**") and seg.endswith("**"):
            _make_run(p, seg[2:-2], bold=True,
                      size_pt=size_pt, font_en=font_en, font_cn=font_cn)
        else:
            _make_run(p, seg, bold=bold,
                      size_pt=size_pt, font_en=font_en, font_cn=font_cn)
    return p


def _add_hr(doc):
    """添加场景分隔线：底边框细线"""
    from docx.shared import Pt
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn

    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after  = Pt(4)

    pPr   = p._p.get_or_add_pPr()
    pBdr  = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"),   "single")
    bottom.set(qn("w:sz"),    "4")       # 0.5pt 细线
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), "999999")  # 浅灰
    pBdr.append(bottom)
    pPr.append(pBdr)


# ─────────────────────── 行分类 ───────────────────────

# 优先级从高到低匹配
_KIND_RULES = [
    (re.compile(r"^# (?!#)"),    "h1"),       # 剧本标题
    (re.compile(r"^### "),       "h3"),        # 集标题
    (re.compile(r"^## "),        "h2"),        # 场景标题
    (re.compile(r"^---\s*$"),    "hr"),        # 分隔线
    (re.compile(r"^△"),          "action"),    # △ 动作/场景描写
    (re.compile(r"^\*\*\S"),     "dialogue"),  # **角色名** 对白
    (re.compile(r"^（"),          "cue"),       # （BGM / 字幕 / 音效）
    (re.compile(r"^【"),          "cue"),       # 【闪回】【闪出】
]


def classify(line: str):
    """返回 (kind, content)；content 对标题类已去掉 # 前缀"""
    s = line.strip()
    if not s:
        return "blank", ""
    for pattern, kind in _KIND_RULES:
        if pattern.match(s):
            content = re.sub(r"^#+\s*", "", s) if kind in ("h1", "h2", "h3") else s
            return kind, content
    return "body", s


# ─────────────────────── 主转换 ───────────────────────

def convert(md_text: str, output_path: str) -> bool:
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    doc = init_document()

    for line in md_text.splitlines():
        kind, content = classify(line)

        if kind == "blank":
            # 间距完全由段落 space_before / space_after 控制，不插入空段落
            continue

        elif kind == "hr":
            _add_hr(doc)

        elif kind == "h1":
            # 剧本总标题：黑体 22pt，居中
            _add_para(doc, content,
                      size_pt=22, font_cn="黑体", font_en="Arial",
                      space_before=0, space_after=12,
                      alignment=WD_ALIGN_PARAGRAPH.CENTER)

        elif kind == "h3":
            # 集标题（### 第N集：...）：黑体 16pt
            _add_para(doc, content,
                      size_pt=16, font_cn="黑体", font_en="Arial",
                      space_before=24, space_after=8)

        elif kind == "h2":
            # 场景标题（## N-M · 地点 · 日夜）：黑体 14pt
            _add_para(doc, content,
                      size_pt=14, font_cn="黑体", font_en="Arial",
                      space_before=18, space_after=6)

        elif kind == "action":
            # △ 动作/场景描写：宋体 12pt 加粗
            _add_para(doc, content,
                      space_before=0, space_after=5)

        elif kind == "dialogue":
            # 对白：**角色名**（括号）台词，内联粗体解析
            _add_para(doc, content,
                      space_before=0, space_after=4)

        elif kind == "cue":
            # 场景指示：（BGM）/ 【闪回】等
            _add_para(doc, content,
                      space_before=0, space_after=4)

        else:  # body
            _add_para(doc, content,
                      space_before=0, space_after=6)

    Path(output_path).parent.mkdir(parents=True, exist_ok=True)
    doc.save(output_path)
    print(f"[完成] Word 文件已生成: {output_path}")
    return True


# ─────────────────────── 入口 ───────────────────────

def main():
    if len(sys.argv) == 2 and sys.argv[1] in {"-h", "--help"}:
        print("用法: python3 export_docx.py <输入.md> <输出.docx>")
        sys.exit(0)

    if len(sys.argv) < 3:
        print("用法: python3 export_docx.py <输入.md> <输出.docx>")
        sys.exit(1)

    if not check_deps():
        print("[提示] Mac:   pip3 install python-docx lxml")
        print("[提示] Win:   pip install python-docx lxml")
        print("[提示] Linux: pip3 install python-docx lxml")
        sys.exit(1)

    input_path  = Path(sys.argv[1])
    output_path = sys.argv[2]
    # sys.argv[3]（旧版 reference-doc 路径）已不再使用，静默忽略

    if not input_path.exists():
        print(f"[错误] 输入文件不存在: {input_path}", file=sys.stderr)
        sys.exit(1)

    md_text = input_path.read_text(encoding="utf-8")
    if not convert(md_text, output_path):
        sys.exit(1)


if __name__ == "__main__":
    main()
